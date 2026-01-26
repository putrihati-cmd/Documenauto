from docx import Document
from docx.shared import Cm, Pt
import sys
import os

def apply_style(target_path, reference_path, output_path):
    """
    Two-Input System:
    1. Reference: Master File (Pemberi Gaya)
    2. Target: Customer File (Penerima Gaya)
    """
    print(f"Applying style from [{reference_path}] to [{target_path}]...")
    
    try:
        # Load Reference
        ref_doc = Document(reference_path)
        
        # Extract Style Rules from Reference (Simple Scanner Logic)
        ref_section = ref_doc.sections[0]
        margin_rules = {
            'top': ref_section.top_margin,
            'bottom': ref_section.bottom_margin,
            'left': ref_section.left_margin,
            'right': ref_section.right_margin,
            'page_width': ref_section.page_width,
            'page_height': ref_section.page_height
        }
        
        # Detect Dominant Font (Mock logic: take first paragraph font or default)
        # In real production, this needs robust scanning.
        font_rules = {'name': 'Times New Roman', 'size': Pt(12)} # Default fallback
        for p in ref_doc.paragraphs[:5]:
            if p.runs and p.runs[0].font.name:
                font_rules['name'] = p.runs[0].font.name
                if p.runs[0].font.size:
                    font_rules['size'] = p.runs[0].font.size
                break

        # Load Target
        target_doc = Document(target_path)
        
        # APPLY MARGINS & PAGE SIZE
        for section in target_doc.sections:
            section.top_margin = margin_rules['top']
            section.bottom_margin = margin_rules['bottom']
            section.left_margin = margin_rules['left']
            section.right_margin = margin_rules['right']
            section.page_width = margin_rules['page_width']
            section.page_height = margin_rules['page_height']
            
        # APPLY FONTS
        for paragraph in target_doc.paragraphs:
            for run in paragraph.runs:
                run.font.name = font_rules['name']
                # Optional: Apply size only if not heading? 
                # For MVP, let's enforce size too to ensure uniformity.
                # run.font.size = font_rules['size'] 

        # Save Result
        target_doc.save(output_path)
        print(f">> Success! Saved to: {output_path}")
        print(f">> Applied Margins: T={margin_rules['top'].cm:.2f}, B={margin_rules['bottom'].cm:.2f}, L={margin_rules['left'].cm:.2f}, R={margin_rules['right'].cm:.2f}")
        print(f">> Applied Font: {font_rules['name']}")
        return True

    except Exception as e:
        print(f"Error applying style: {e}")
        return False

if __name__ == "__main__":
    if len(sys.argv) < 4:
        print("Usage: python style_applicator.py <target_file> <reference_category_or_file> <output_file>")
        sys.exit(1)
        
    target = sys.argv[1]
    reference = sys.argv[2] # Can be a path or a category name
    output = sys.argv[3]
    
    # If reference is just a category name (e.g. 'skripsi'), map it to the master file in templates
    # For this POC, we assume absolute path is passed or we resolve it here.
    # Simple resolution logic:
    if not os.path.exists(reference):
        # Check in templates folder
        possible_path = f"/app/templates/{reference}/master.docx" # Docker path
        if os.path.exists(possible_path):
            reference = possible_path
        else:
             # Fallback check for local dev environment
            possible_path_local = f"templates/{reference}/master.docx"
            if os.path.exists(possible_path_local):
                reference = possible_path_local
    
    apply_style(target, reference, output)
