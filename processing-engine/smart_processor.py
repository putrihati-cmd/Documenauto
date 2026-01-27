from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import sys
import re

class SmartProcessor:
    def __init__(self, rules_json):
        """Initialize processor with rules"""
        self.rules = json.loads(rules_json) if isinstance(rules_json, str) else rules_json
        self.warnings = []
    
    def process_document(self, input_path, output_path):
        """Main processing function"""
        print(f"[SmartProcessor] Loading document: {input_path}")
        doc = Document(input_path)
        
        # Apply global rules
        print("[SmartProcessor] Applying global rules...")
        self.apply_global_rules(doc)
        
        # Apply section-specific rules
        print("[SmartProcessor] Applying section-specific rules...")
        self.apply_section_rules(doc)
        
        # Save
        doc.save(output_path)
        print(f"[SmartProcessor] Document saved: {output_path}")
        
        # Print warnings
        if self.warnings:
            print("\n⚠️  WARNINGS:")
            for warning in self.warnings:
                print(f"  - {warning}")
        
        return {
            'success': True,
            'warnings': self.warnings,
            'output_path': output_path
        }
    
    def apply_global_rules(self, doc):
        """Apply global formatting rules"""
        if 'global' not in self.rules:
            return
        
        global_rules = self.rules['global']
        
        # Apply margins
        if 'margins' in global_rules:
            margins = global_rules['margins']
            for section in doc.sections:
                if 'top' in margins:
                    section.top_margin = self._parse_size(margins['top'])
                if 'bottom' in margins:
                    section.bottom_margin = self._parse_size(margins['bottom'])
                if 'left' in margins:
                    section.left_margin = self._parse_size(margins['left'])
                if 'right' in margins:
                    section.right_margin = self._parse_size(margins['right'])
        
        # Apply font and spacing to all paragraphs
        if 'font' in global_rules or 'line_spacing' in global_rules:
            for para in doc.paragraphs:
                # Skip if paragraph has specific section rule
                if not self._has_section_marker(para):
                    self._apply_paragraph_format(para, global_rules)
    
    def apply_section_rules(self, doc):
        """Apply section-specific rules"""
        if 'sections' not in self.rules:
            return
        
        sections = self.rules['sections']
        
        # Detect and apply abstract rules
        if 'abstract' in sections:
            self.apply_abstract_rules(doc, sections['abstract'])
        
        # Detect and apply chapter heading rules
        if 'chapter_headings' in sections:
            self.apply_chapter_rules(doc, sections['chapter_headings'])
        
        # Detect and apply bibliography rules
        if 'bibliography' in sections:
            self.apply_bibliography_rules(doc, sections['bibliography'])
        
        # Detect and apply table of contents rules
        if 'table_of_contents' in sections:
            self.apply_toc_rules(doc, sections['table_of_contents'])
    
    def apply_abstract_rules(self, doc, rules):
        """Apply formatting to abstract section"""
        abstract_paras = self._find_section(doc, ['abstract', 'abstrak'])
        
        if not abstract_paras:
            print("  [Abstract] Section not found")
            return
        
        print(f"  [Abstract] Found {len(abstract_paras)} paragraphs")
        
        # Apply formatting
        for para in abstract_paras:
            self._apply_paragraph_format(para, rules)
        
        # Check word count
        if 'max_words' in rules:
            total_words = sum(len(p.text.split()) for p in abstract_paras)
            if total_words > rules['max_words']:
                self.warnings.append(
                    f"Abstract too long: {total_words} words (max: {rules['max_words']})"
                )
    
    def apply_chapter_rules(self, doc, rules):
        """Apply formatting to chapter headings (BAB)"""
        chapter_paras = self._find_chapters(doc)
        
        if not chapter_paras:
            print("  [Chapters] No chapters found")
            return
        
        print(f"  [Chapters] Found {len(chapter_paras)} chapters")
        
        for para in chapter_paras:
            self._apply_paragraph_format(para, rules)
            
            # Apply text transform
            if rules.get('text_transform') == 'uppercase':
                for run in para.runs:
                    run.text = run.text.upper()
    
    def apply_bibliography_rules(self, doc, rules):
        """Apply formatting to bibliography section"""
        biblio_paras = self._find_section(doc, [
            'daftar pustaka', 'bibliography', 'references', 'referensi'
        ])
        
        if not biblio_paras:
            print("  [Bibliography] Section not found")
            return
        
        print(f"  [Bibliography] Found {len(biblio_paras)} entries")
        
        for para in biblio_paras:
            self._apply_paragraph_format(para, rules)
            
            # Apply hanging indent
            if 'indent_hanging' in rules:
                indent_size = self._parse_size(rules['indent_hanging'])
                para.paragraph_format.left_indent = indent_size
                para.paragraph_format.first_line_indent = -indent_size
    
    def apply_toc_rules(self, doc, rules):
        """Apply formatting to table of contents"""
        toc_paras = self._find_section(doc, [
            'daftar isi', 'table of contents', 'contents'
        ])
        
        if not toc_paras:
            print("  [TOC] Section not found")
            return
        
        print(f"  [TOC] Found {len(toc_paras)} entries")
        
        for para in toc_paras:
            self._apply_paragraph_format(para, rules)
    
    def _find_section(self, doc, keywords):
        """Find paragraphs in a section by keywords"""
        found = False
        section_paras = []
        
        for para in doc.paragraphs:
            text_lower = para.text.lower().strip()
            
            # Check if this is the section start
            if not found:
                for keyword in keywords:
                    if keyword in text_lower:
                        found = True
                        break
            
            # If found, collect paragraphs until next major section
            elif found:
                # Stop at next BAB or major section
                if re.match(r'^(bab|chapter)\s+[ivx\d]+', text_lower):
                    break
                if para.text.strip():
                    section_paras.append(para)
        
        return section_paras
    
    def _find_chapters(self, doc):
        """Find all chapter headings (BAB)"""
        chapters = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            # Match "BAB I", "BAB II", "CHAPTER 1", etc.
            if re.match(r'^(BAB|CHAPTER)\s+[IVX\d]+', text, re.IGNORECASE):
                chapters.append(para)
        
        return chapters
    
    def _apply_paragraph_format(self, para, rules):
        """Apply formatting rules to a paragraph"""
        # Apply to all runs in paragraph
        for run in para.runs:
            # Font name
            if 'font' in rules and 'name' in rules['font']:
                run.font.name = rules['font']['name']
            
            # Font size
            if 'font_size' in rules:
                run.font.size = self._parse_font_size(rules['font_size'])
            elif 'font' in rules and 'size' in rules['font']:
                run.font.size = self._parse_font_size(rules['font']['size'])
            
            # Font weight
            if rules.get('font_weight') == 'bold':
                run.font.bold = True
            
            # Font style
            if rules.get('title_style') == 'italic':
                run.font.italic = True
        
        # Line spacing
        if 'line_spacing' in rules:
            spacing_value = float(rules['line_spacing'])
            para.paragraph_format.line_spacing = spacing_value
        
        # Spacing before/after
        if 'spacing_before' in rules:
            para.paragraph_format.space_before = self._parse_size(rules['spacing_before'])
        if 'spacing_after' in rules:
            para.paragraph_format.space_after = self._parse_size(rules['spacing_after'])
    
    def _parse_size(self, size_str):
        """Parse size string to docx units"""
        if isinstance(size_str, (int, float)):
            return Cm(size_str)
        
        size_str = str(size_str).lower().strip()
        
        if 'cm' in size_str:
            value = float(size_str.replace('cm', '').strip())
            return Cm(value)
        elif 'in' in size_str:
            value = float(size_str.replace('in', '').strip())
            return Inches(value)
        elif 'pt' in size_str:
            value = float(size_str.replace('pt', '').strip())
            return Pt(value)
        else:
            # Default to cm
            return Cm(float(size_str))
    
    def _parse_font_size(self, size_str):
        """Parse font size to points"""
        if isinstance(size_str, (int, float)):
            return Pt(size_str)
        
        size_str = str(size_str).lower().strip()
        value = float(size_str.replace('pt', '').strip())
        return Pt(value)
    
    def _has_section_marker(self, para):
        """Check if paragraph is part of a special section"""
        text = para.text.lower().strip()
        markers = [
            'abstract', 'abstrak', 'bab', 'chapter',
            'daftar pustaka', 'bibliography', 'daftar isi'
        ]
        return any(marker in text for marker in markers)


def main():
    if len(sys.argv) < 4:
        print("Usage: python smart_processor.py <input.docx> <rules.json> <output.docx>")
        sys.exit(1)
    
    input_file = sys.argv[1]
    rules_file = sys.argv[2]
    output_file = sys.argv[3]
    
    # Load rules
    with open(rules_file, 'r') as f:
        rules = json.load(f)
    
    # Process
    processor = SmartProcessor(rules)
    result = processor.process_document(input_file, output_file)
    
    # Print result
    print("\n✅ Processing complete!")
    print(json.dumps(result, indent=2))


if __name__ == "__main__":
    main()
