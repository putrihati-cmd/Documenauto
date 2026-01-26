import os
from docx import Document
from docx.shared import Cm, Pt

# 1. Buat File Dummy (Pura-pura ini Skripsi User)
def create_dummy_skripsi():
    doc = Document()
    
    # Bikin Margin Salah (misal 2cm semua)
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Tambah Judul
    p = doc.add_paragraph()
    runner = p.add_run('SKRIPSI TENTANG AI')
    runner.bold = True
    runner.font.name = 'Arial' # Salah font (harusnya TNR)
    runner.font.size = Pt(14)

    # Tambah Isi
    doc.add_paragraph('Bab I: Pendahuluan...')
    
    doc.save('sample_skripsi.docx')
    print("File 'sample_skripsi.docx' dibuat dengan format yang SALAH (Margin 2cm, Font Arial).")

# 2. Analisa File (Kemampuan "Memahami")
def analyze_structure(filename):
    print(f"\n--- Menganalisa {filename} ---")
    doc = Document(filename)
    
    # Cek Margin
    section = doc.sections[0]
    top = round(section.top_margin.cm, 1)
    bottom = round(section.bottom_margin.cm, 1)
    left = round(section.left_margin.cm, 1)
    right = round(section.right_margin.cm, 1)
    
    print(f"[Check Margin] Terdeteksi: T={top}, B={bottom}, L={left}, R={right}")
    if left == 4.0 and top == 4.0:
        print(">> Status: Margin Sesuai Standar Skripsi.")
    else:
        print(f">> Status: TIDAK SESUAI (Standar biasanya 4,4,3,3).")

    # Cek Font
    font_name = "Unknown"
    for p in doc.paragraphs:
        if p.runs:
            font_name = p.runs[0].font.name
            break # Cek paragraf pertama aja
            
    print(f"[Check Font]   Terdeteksi: {font_name}")
    if font_name == 'Times New Roman':
        print(">> Status: Font Sesuai.")
    else:
        print(">> Status: TIDAK SESUAI (Harusnya Times New Roman).")

if __name__ == "__main__":
    create_dummy_skripsi()
    analyze_structure('sample_skripsi.docx')
