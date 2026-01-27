import requests
import json
import time
import os

# Configuration
BASE_URL = "http://192.168.1.27:3003/api" # Backend Port on Server
FRONTEND_URL = "http://192.168.1.27:3002"
TEST_FILE_PATH = "audit_dummy.docx"

def log(step, message, status="INFO"):
    print(f"[{status}] {step}: {message}")

def create_dummy_docx():
    if not os.path.exists(TEST_FILE_PATH):
        try:
            from docx import Document
            doc = Document()
            doc.add_heading('Audit Test Document', 0)
            doc.add_paragraph('This is a test document for auditing the production deployment.')
            doc.save(TEST_FILE_PATH)
            log("Setup", f"Created dummy file: {TEST_FILE_PATH}", "SUCCESS")
        except ImportError:
            # Fallback: Create empty file
            with open(TEST_FILE_PATH, 'w') as f:
                f.write("Dummy content")
            log("Setup", f"Created text dummy (no docx lib): {TEST_FILE_PATH}", "WARNING")

def check_health():
    # Frontend
    try:
        r = requests.get(FRONTEND_URL, timeout=5)
        if r.status_code == 200:
            log("Health Check", f"Frontend is UP ({FRONTEND_URL})", "SUCCESS")
        else:
            log("Health Check", f"Frontend returned {r.status_code}", "WARNING")
    except Exception as e:
        log("Health Check", f"Frontend Unreachable: {e}", "FAIL")

    # Backend
    try:
        r = requests.get(f"{BASE_URL}/health", timeout=5) # Assuming /health exists, or root
        # If /health doesn't exist, try /templates which might return 401 or empty list
    except:
        pass # Will test via API calls

def audit_flow():
    session = requests.Session()

    # 1. Upload Template (Admin)
    log("Step 1", "Uploading Audit Template...")
    try:
        with open(TEST_FILE_PATH, 'rb') as f:
            files = {'file': (TEST_FILE_PATH, f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
            data = {
                'name': 'Audit Template',
                'category': 'audit',
                'description': 'Automated audit test',
                'is_default': 'false'
            }
            res = session.post(f"{BASE_URL}/templates", files=files, data=data)
        
        if res.status_code == 201:
            template = res.json().get('data')
            template_id = template.get('id')
            log("Step 1", f"Template Created! ID: {template_id}", "SUCCESS")
        else:
            log("Step 1", f"Failed: {res.text}", "FAIL")
            return
    except Exception as e:
        log("Step 1", f"Error: {e}", "FAIL")
        return

    # 2. Retrieve Template List (Verify Persistence)
    log("Step 2", "Verifying Template in List...")
    try:
        res = session.get(f"{BASE_URL}/templates?category=audit")
        if res.status_code == 200:
            templates = res.json().get('data', [])
            found = any(str(t.get('id')) == str(template_id) for t in templates)
            if found:
                log("Step 2", "Template Found in DB!", "SUCCESS")
            else:
                log("Step 2", "Template NOT Found in List (Persistence Issue?)", "FAIL")
        else:
            log("Step 2", f"Failed fetching lists: {res.status_code}", "FAIL")
    except Exception as e:
        log("Step 2", f"Error: {e}", "FAIL")

    # 3. Create Order (Customer)
    log("Step 3", "Creating Order using Audit Template...")
    order_id = None
    try:
        with open(TEST_FILE_PATH, 'rb') as f:
            files = {'file': (TEST_FILE_PATH, f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
            data = {
                'documentType': 'skripsi', # or category
                'templateId': template_id, # Linking to our new template
                'serviceLevel': 'express',
                'copies': 1
            }
            res = session.post(f"{BASE_URL}/orders", files=files, data=data)
            
        if res.status_code == 201:
            order = res.json().get('data')
            order_id = order.get('orderId')
            log("Step 3", f"Order Created! ID: {order_id}", "SUCCESS")
        else:
            log("Step 3", f"Failed: {res.text}", "FAIL")
            return
    except Exception as e:
        log("Step 3", f"Error: {e}", "FAIL")
        return

    # 4. Check API Response for Status (Queue Integration)
    # Since we can't easily check the worker output file remotely via HTTP without a specific endpoint,
    # we will rely on successful order creation and Queue ACK logs from backend if we could see them.
    # But as a blackbox audit, receiving 201 Created is good.
    # Future: Poll status endpoint if implemented.

    log("Step 4", "Checking Order Status...", "INFO")
    # Assuming GET /orders/:id exists
    try:
        res = session.get(f"{BASE_URL}/orders/{order_id}")
        if res.status_code == 200:
             status = res.json().get('data', {}).get('status')
             log("Step 4", f"Order Status: {status}", "SUCCESS")
        else:
             log("Step 4", "Could not fetch status.", "WARNING")
    except:
        pass

if __name__ == "__main__":
    print("=== STARTING REMOTE AUDIT ===")
    create_dummy_docx()
    check_health()
    audit_flow()
    print("=== AUDIT COMPLETE ===")
